import streamlit as st
import pandas as pd
import gspread
from oauth2client.service_account import ServiceAccountCredentials
import time
from datetime import datetime, date, timedelta
import requests
import base64
from io import BytesIO

# --- KIỂM TRA THƯ VIỆN WORD ---
try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    st.error("⚠️ Lỗi: Chưa cài thư viện python-docx. Vui lòng chạy lệnh: pip install python-docx")
    st.stop()

# --- CẤU HÌNH HỆ THỐNG ---
st.set_page_config(page_title="HR System Pro", layout="wide", page_icon="💎")

# ĐỊNH NGHĨA QUY TRÌNH & DEADLINE (SLA)
WORKFLOW = {
    "Mới nhận": {"step": 1, "sla": 2},      
    "Sơ loại":  {"step": 2, "sla": 3},
    "Phỏng vấn": {"step": 3, "sla": 5},
    "Chờ kết quả": {"step": 4, "sla": 7},
    "Đạt / Chờ đi làm": {"step": 5, "sla": 10},
    "Đã đi làm": {"step": 6, "sla": 0},     
    "Loại": {"step": 6, "sla": 0},           
    "Nghỉ việc": {"step": 6, "sla": 0}
}

# Link Apps Script (Giữ nguyên)
APPS_SCRIPT_URL = "https://script.google.com/macros/s/AKfycbzKueqCnPonJ1MsFzQpQDk7ihgnVVQyNHMUyc_dx6AocsDu1jW1zf6Gr9VgqMD4D00/exec"

# --- CSS GIAO DIỆN ---
st.markdown("""
    <style>
    [data-testid="stSidebar"] {background-color: #f8f9fa;}
    [data-testid="stSidebar"] .stButton > button {
        width: 100%; height: 50px; border: none; border-radius: 8px;
        background-color: white; color: #333; font-weight: 600;
        box-shadow: 0 1px 3px rgba(0,0,0,0.1); text-align: left; padding-left: 15px; margin-bottom: 5px;
    }
    [data-testid="stSidebar"] .stButton > button:hover {background-color: #e3f2fd; color: #1565c0;}
    
    .sticky-note {
        background-color: #fff9c4; padding: 10px; border-radius: 5px; 
        border-left: 5px solid #fbc02d; font-size: 14px;
        font-family: 'Arial', sans-serif; margin-bottom: 10px; color: #333;
    }
    .note-badge {
        background-color: #fbc02d; color: black; padding: 2px 8px; 
        border-radius: 10px; font-size: 0.8em; font-weight: bold; margin-left: 10px;
    }
    </style>
""", unsafe_allow_html=True)

# --- KẾT NỐI GOOGLE SHEETS ---
@st.cache_resource
def get_gcp_service():
    try:
        scope = ["https://spreadsheets.google.com/feeds", "https://www.googleapis.com/auth/drive"]
        creds = ServiceAccountCredentials.from_json_keyfile_dict(st.secrets["gcp_service_account"], scope)
        client = gspread.authorize(creds)
        return client
    except: return None

client = get_gcp_service()
if not client: st.error("⚠️ Lỗi kết nối Secrets!"); st.stop()

try:
    sheet_ungvien = client.open("TuyenDungKCN_Data").worksheet("UngVien")
    sheet_users = client.open("TuyenDungKCN_Data").worksheet("Users")
except: st.error("⚠️ Không tìm thấy file Excel hoặc Sheet UngVien/Users."); st.stop()

# --- CÁC HÀM HỖ TRỢ ---
def upload_via_appsscript(file_obj, file_name):
    try:
        file_bytes = file_obj.getvalue()
        base64_str = base64.b64encode(file_bytes).decode('utf-8')
        payload = {"base64": base64_str, "filename": file_name, "mimeType": file_obj.type}
        response = requests.post(APPS_SCRIPT_URL, json=payload)
        if response.status_code == 200:
            res_json = response.json()
            if res_json.get("result") == "success": return res_json.get("link")
    except: pass
    return None

def convert_drive_link(link):
    if "id=" in link:
        file_id = link.split("id=")[1]
        return f"https://drive.google.com/thumbnail?id={file_id}&sz=w1000" 
    return link

def parse_date_vn(date_str):
    """Chuyển đổi chuỗi ngày VN sang đối tượng Date"""
    try:
        return datetime.strptime(date_str, "%d/%m/%Y").date()
    except:
        return date(2000, 1, 1) # Mặc định nếu lỗi

def calculate_deadline_status(start_date_str, status):
    try:
        if status not in WORKFLOW or WORKFLOW[status]['sla'] == 0:
            return None, "completed"
        start_date = datetime.strptime(start_date_str, "%d/%m/%Y")
        deadline_date = start_date + timedelta(days=WORKFLOW[status]['sla'])
        days_left = (deadline_date - datetime.now()).days
        return days_left, deadline_date.strftime("%d/%m/%Y")
    except: return None, None

def create_word_file(data):
    doc = Document()
    style = doc.styles['Normal']; font = style.font; font.name = 'Times New Roman'; font.size = Pt(13)
    
    head = doc.add_heading(f"HỒ SƠ ỨNG VIÊN: {data['HoTen']}", 0)
    head.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in head.runs:
        run.font.name = 'Times New Roman'; run.font.size = Pt(16); run.font.bold = True; run.font.color.rgb = RGBColor(0, 0, 0)

    def add_line(label, value):
        p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(6)
        runner = p.add_run(f"{label}: "); runner.font.name = 'Times New Roman'; runner.font.bold = True
        runner_val = p.add_run(str(value) if value else ""); runner_val.font.name = 'Times New Roman'

    p_sub = doc.add_paragraph(); p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = p_sub.add_run(f"(Vị trí: {data['ViTri']} | Trạng thái: {data['TrangThai']})"); run_sub.font.name = 'Times New Roman'; run_sub.italic = True
    doc.add_paragraph("")

    h1 = doc.add_heading('I. THÔNG TIN CÁ NHÂN', level=1)
    for run in h1.runs: run.font.name = 'Times New Roman'; run.font.size = Pt(14); run.font.color.rgb = RGBColor(0,0,0)
    add_line("Họ và tên", data['HoTen']); add_line("Ngày sinh", data['NamSinh'])
    add_line("Số điện thoại", data['SDT']); add_line("CCCD", data.get('CCCD', ''))
    add_line("Quê quán", data['QueQuan'])

    h2 = doc.add_heading('II. THÔNG TIN KHÁC', level=1)
    for run in h2.runs: run.font.name = 'Times New Roman'; run.font.size = Pt(14); run.font.color.rgb = RGBColor(0,0,0)
    add_line("Nguồn tuyển dụng", data.get('Nguồn', '')); add_line("Đăng ký xe tuyến", data.get('XeTuyen', ''))
    add_line("Nhu cầu KTX", data.get('KTX', '')); add_line("Ghi chú hiện tại", data.get('GhiChu', ''))

    doc.add_paragraph("")
    p_footer = doc.add_paragraph(f"Ngày xuất hồ sơ: {datetime.now().strftime('%d/%m/%Y')}"); p_footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in p_footer.runs: run.font.name = 'Times New Roman'; run.font.italic = True; run.font.size = Pt(11)

    buf = BytesIO(); doc.save(buf); buf.seek(0)
    return buf

# --- SESSION ---
if 'logged_in' not in st.session_state: st.session_state.logged_in = False
if 'current_page' not in st.session_state: st.session_state.current_page = "dashboard"
def set_page(page_name): st.session_state.current_page = page_name

# --- LOGIN ---
def login_screen():
    st.markdown("<br><h1 style='text-align: center; color:#1565c0'>🔐 HR SYSTEM PRO</h1>", unsafe_allow_html=True)
    c1, c2, c3 = st.columns([1,1,1])
    with c2:
        tab1, tab2 = st.tabs(["ĐĂNG NHẬP", "ĐĂNG KÝ"])
        with tab1:
            with st.form("login"):
                u = st.text_input("Username"); p = st.text_input("Password", type="password")
                if st.form_submit_button("VÀO HỆ THỐNG", use_container_width=True):
                    users = sheet_users.get_all_records()
                    for user in users:
                        if str(user['Username']) == u and str(user['Password']) == p:
                            st.session_state.logged_in = True; st.session_state.user_role = user['Role']
                            st.session_state.user_name = user['HoTen']; st.rerun()
                    st.error("Sai thông tin!")
        with tab2:
            with st.form("reg"):
                nu = st.text_input("User mới"); np = st.text_input("Pass mới", type="password"); nn = st.text_input("Họ tên")
                if st.form_submit_button("TẠO TÀI KHOẢN"):
                    existing = sheet_users.col_values(1)
                    if nu in existing: st.warning("Tên tồn tại!")
                    else: sheet_users.append_row([nu, np, "staff", nn]); st.success("OK! Mời đăng nhập.")

# --- MAIN APP ---
def main_app():
    raw_data = sheet_ungvien.get_all_records()
    df = pd.DataFrame(raw_data)
    df.columns = [c.strip() for c in df.columns]

    with st.sidebar:
        st.markdown(f"### 👤 {st.session_state.user_name}")
        st.caption(f"Role: {st.session_state.user_role.upper()}")
        st.markdown("---")
        if st.button("🏠 DASHBOARD"): set_page("dashboard")
        if st.button("📝 NHẬP HỒ SƠ"): set_page("input")
        if st.button("🔍 DANH SÁCH"): set_page("list")
        if st.session_state.user_role == "admin":
            st.markdown("---"); 
            if st.button("⚙️ QUẢN TRỊ"): set_page("admin")
        st.markdown("<br>", unsafe_allow_html=True)
        if st.button("🚪 Đăng xuất"): st.session_state.logged_in = False; st.rerun()

    # 1. DASHBOARD
    if st.session_state.current_page == "dashboard":
        st.title("📊 Tổng Quan Tuyển Dụng")
        if not df.empty:
            c1, c2, c3, c4 = st.columns(4)
            tt_col = 'TrangThai' if 'TrangThai' in df.columns else None
            
            with c1: st.metric("Tổng Hồ Sơ", len(df), delta=f"+{len(df[df[tt_col]=='Mới nhận'])} mới")
            with c2: st.metric("Đã Đi Làm", len(df[df[tt_col]=='Đã đi làm']) if tt_col else 0)
            with c3: st.metric("Phỏng Vấn", len(df[df[tt_col]=='Phỏng vấn']) if tt_col else 0)
            
            overdue_count = 0
            if 'NgayNhap' in df.columns and tt_col:
                for _, row in df.iterrows():
                    days_left, _ = calculate_deadline_status(row['NgayNhap'], row[tt_col])
                    if days_left is not None and days_left < 0: overdue_count += 1
            with c4: st.metric("⚠️ Quá Hạn", overdue_count, delta_color="inverse")
            
            st.markdown("---")
            col_chart, col_table = st.columns([1, 1])
            with col_chart:
                st.subheader("Tiến độ")
                if tt_col: st.bar_chart(df[tt_col].value_counts())
            with col_table:
                st.subheader("Top Tuyển Dụng")
                if 'NguoiTuyen' in df.columns:
                    top = df['NguoiTuyen'].value_counts().reset_index(); top.columns = ['Recruiter', 'Count']
                    st.dataframe(top, use_container_width=True, hide_index=True)

    # 2. NHẬP LIỆU
    elif st.session_state.current_page == "input":
        st.header("📝 Nhập Hồ Sơ Mới")
        with st.form("input_form"):
            col_img, col_info = st.columns([1, 3])
            with col_img: uploaded_file = st.file_uploader("Upload ảnh (3x4)", type=['jpg','png','jpeg'])
            with col_info:
                name = st.text_input("Họ tên (*)")
                phone = st.text_input("SĐT (*)")
                cccd = st.text_input("CCCD")
            
            r1, r2, r3 = st.columns(3)
            dob = r1.date_input("Ngày sinh", value=date(2000, 1, 1), min_value=date(1960, 1, 1))
            hometown = r2.text_input("Quê quán")
            pos = r3.selectbox("Vị trí", ["Công nhân", "Kỹ thuật", "Kho", "Bảo vệ", "Tạp vụ", "Khác"])
            
            r4, r5 = st.columns(2)
            source = r4.selectbox("Nguồn", ["Facebook", "Zalo", "Trực tiếp", "Người giới thiệu"])
            note = st.text_area("Ghi chú ban đầu")
            
            st.markdown("---")
            fb = st.text_input("Link Facebook"); tt = st.text_input("Link TikTok")
            r6, r7, r8 = st.columns(3)
            bus = r6.selectbox("Xe tuyến", ["Tự túc", "Tuyến A", "Tuyến B"])
            doc = r7.selectbox("Giấy tờ", ["Chưa có", "Đủ giấy tờ"])
            ktx = r8.selectbox("Ký túc xá", ["Không", "Có"])

            if st.form_submit_button("🚀 LƯU HỒ SƠ", type="primary"):
                if name and phone:
                    with st.spinner("Đang xử lý..."):
                        link_drive = ""
                        if uploaded_file: link_drive = upload_via_appsscript(uploaded_file, f"{name}_{phone}.jpg")
                        
                        now_str = datetime.now().strftime("%d/%m/%Y %H:%M")
                        history_log = f"[{now_str}] {st.session_state.user_name}: Tạo mới hồ sơ."
                        
                        row = [datetime.now().strftime("%d/%m/%Y"), name.upper(), dob.strftime("%d/%m/%Y"), hometown, 
                               f"'{phone}", f"'{cccd}", pos, "Mới nhận", note, source, link_drive, bus, ktx, 
                               st.session_state.user_name, fb, tt, doc, history_log]
                        sheet_ungvien.append_row(row)
                        st.success("✅ Đã thêm hồ sơ!"); time.sleep(1); st.rerun()
                else: st.error("Thiếu Tên hoặc SĐT!")

    # 3. DANH SÁCH (TÍNH NĂNG FULL)
    elif st.session_state.current_page == "list":
        st.header("🗂️ Quản Lý Hồ Sơ")
        if st.button("🔄 Cập nhật dữ liệu"): st.cache_data.clear(); st.rerun()

        if not df.empty:
            search = st.text_input("🔎 Tìm kiếm:")
            df_show = df[df.astype(str).apply(lambda x: x.str.contains(search, case=False)).any(axis=1)] if search else df
            
            # Filter
            st_filter = st.multiselect("Lọc trạng thái", list(WORKFLOW.keys()))
            if st_filter: df_show = df_show[df_show['TrangThai'].isin(st_filter)]

            st.dataframe(df_show[['HoTen', 'SDT', 'ViTri', 'TrangThai']], use_container_width=True, hide_index=True)
            st.markdown("---")

            for i, row in df_show.iterrows():
                with st.container(border=True):
                    # --- Header ---
                    c1, c2 = st.columns([1, 4])
                    with c1:
                        raw_link = str(row.get('LinkAnh', ''))
                        st.image(convert_drive_link(raw_link) if "http" in raw_link else "https://via.placeholder.com/150", width=100)
                    with c2:
                        # TÍNH NĂNG 1: HIỂN THỊ NOTE CẠNH TÊN
                        note_content = row.get('GhiChu', '')
                        note_html = f'<span class="note-badge">📌 {note_content}</span>' if note_content else ""
                        st.markdown(f"### {row['HoTen']} ({row.get('NamSinh', '')}) {note_html}", unsafe_allow_html=True)
                        
                        # Deadline
                        days_left, deadline_date = calculate_deadline_status(row['NgayNhap'], row['TrangThai'])
                        sla_txt = f" | Deadline: {deadline_date}" if days_left is not None else ""
                        st.markdown(f"**{row['ViTri']}** | `{row['TrangThai']}`{sla_txt}")
                        
                        # Progress
                        cur_step = WORKFLOW.get(row['TrangThai'], {}).get('step', 0)
                        st.progress(cur_step / 6, text=f"Tiến độ: Bước {cur_step}/6")

                    # --- Tabs ---
                    t1, t2, t3 = st.tabs(["ℹ️ Chi Tiết", "📝 Ghi Chú & Lịch Sử", "⚙️ Chỉnh Sửa & Tác Vụ"])
                    
                    with t1:
                        c_a, c_b = st.columns(2)
                        with c_a:
                            st.write(f"📞 SĐT: {row['SDT']}"); st.write(f"🆔 CCCD: {row.get('CCCD', '--')}")
                            st.write(f"🏠 Quê: {row['QueQuan']}")
                        with c_b:
                            st.write(f"🚌 Xe: {row.get('XeTuyen', '--')}"); st.write(f"🏨 KTX: {row.get('KTX', '--')}")
                            st.write(f"📄 Giấy tờ: {row.get('GiayTo', '--')}")
                        
                        if st.button("📄 Tải File Word", key=f"dl_{i}"):
                             doc_file = create_word_file(row)
                             st.download_button("Click tải xuống", doc_file, f"{row['HoTen']}.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")

                    with t2:
                        st.write("#### Sticky Note hiện tại:")
                        if row.get('GhiChu'):
                            st.markdown(f"""<div class="sticky-note">📝 {row['GhiChu']}</div>""", unsafe_allow_html=True)
                        else:
                            st.info("Chưa có ghi chú nào.")
                            
                        with st.expander("📜 Xem lịch sử ghi chú & thay đổi"):
                            st.markdown(str(row.get('LichSu', '')).replace('\n', '<br>'), unsafe_allow_html=True)

                    # --- TAB 3: CHỈNH SỬA TOÀN BỘ (SỬA ĐƯỢC CẢ SĐT & NGÀY SINH) ---
                    with t3:
                        with st.form(key=f"full_edit_{i}"):
                            st.write("#### ✏️ Cập nhật thông tin hồ sơ")
                            # 1. Thông tin cá nhân
                            ec1, ec2, ec3 = st.columns(3)
                            new_name = ec1.text_input("Họ tên", value=row['HoTen'])
                            
                            # TÍNH NĂNG 2: SỬA NGÀY SINH & SĐT
                            current_dob = parse_date_vn(str(row.get('NamSinh', '')))
                            new_dob = ec2.date_input("Ngày sinh", value=current_dob)
                            
                            # Lưu ý: SĐT dùng để tìm dòng, nên cần cẩn thận khi sửa
                            new_phone = ec3.text_input("SĐT (Cẩn thận khi sửa)", value=str(row['SDT']).replace("'", ""))

                            ec4, ec5 = st.columns(2)
                            curr_cccd = str(row.get('CCCD','')).replace("'","")
                            new_cccd = ec4.text_input("CCCD", value=curr_cccd)
                            new_hometown = ec5.text_input("Quê quán", value=row['QueQuan'])
                            
                            # 2. Công việc
                            ec6, ec7 = st.columns(2)
                            pos_opts = ["Công nhân", "Kỹ thuật", "Kho", "Bảo vệ", "Tạp vụ", "Khác"]
                            p_idx = pos_opts.index(row['ViTri']) if row['ViTri'] in pos_opts else 0
                            new_pos = ec6.selectbox("Vị trí", pos_opts, index=p_idx)

                            wf_keys = list(WORKFLOW.keys())
                            s_idx = wf_keys.index(row['TrangThai']) if row['TrangThai'] in wf_keys else 0
                            new_status = ec7.selectbox("Trạng thái", wf_keys, index=s_idx)

                            st.markdown("---")
                            # 3. Ghi chú
                            st.write("#### 📝 Ghi chú (Sticky Note)")
                            new_note = st.text_area("Nội dung ghi chú sẽ hiển thị cạnh tên:", value=row.get('GhiChu', ''))

                            # Nút lưu duy nhất
                            if st.form_submit_button("💾 LƯU TẤT CẢ THAY ĐỔI"):
                                try:
                                    # Tìm dòng bằng SĐT CŨ (quan trọng)
                                    cell = sheet_ungvien.find(str(row['SDT']))
                                    if cell:
                                        # Tạo log lịch sử
                                        now = datetime.now().strftime("%d/%m/%Y %H:%M")
                                        log_entry = ""
                                        
                                        # Kiểm tra thay đổi để ghi log
                                        if new_status != row['TrangThai']:
                                            log_entry += f"[{now}] {st.session_state.user_name}: Đổi trạng thái '{row['TrangThai']}' -> '{new_status}'\n"
                                        
                                        # TÍNH NĂNG 3: LOG LỊCH SỬ GHI CHÚ
                                        if new_note != row.get('GhiChu', ''):
                                            old_n = row.get('GhiChu', '(Trống)')
                                            log_entry += f"[{now}] {st.session_state.user_name}: Sửa ghi chú: '{old_n}' -> '{new_note}'\n"

                                        if new_phone != str(row['SDT']).replace("'", ""):
                                             log_entry += f"[{now}] {st.session_state.user_name}: Đổi SĐT từ {row['SDT']} -> {new_phone}\n"

                                        # CẬP NHẬT GOOGLE SHEET (Mapping đúng cột)
                                        # Giả định thứ tự: [Ngay, HoTen, NamSinh, Que, SDT, CCCD, ViTri, TrangThai, GhiChu...]
                                        sheet_ungvien.update_cell(cell.row, 2, new_name.upper()) # Col 2: Name
                                        sheet_ungvien.update_cell(cell.row, 3, new_dob.strftime("%d/%m/%Y")) # Col 3: DOB (Mới)
                                        sheet_ungvien.update_cell(cell.row, 4, new_hometown)     # Col 4: Que
                                        sheet_ungvien.update_cell(cell.row, 5, f"'{new_phone}")  # Col 5: SDT (Mới)
                                        sheet_ungvien.update_cell(cell.row, 6, f"'{new_cccd}")   # Col 6: CCCD
                                        sheet_ungvien.update_cell(cell.row, 7, new_pos)          # Col 7: Pos
                                        sheet_ungvien.update_cell(cell.row, 8, new_status)       # Col 8: Status
                                        sheet_ungvien.update_cell(cell.row, 9, new_note)         # Col 9: Note
                                        
                                        if log_entry:
                                            old_hist = row.get('LichSu', '')
                                            sheet_ungvien.update_cell(cell.row, 18, log_entry + str(old_hist))

                                        st.success("✅ Đã cập nhật thành công!"); time.sleep(1); st.rerun()
                                    else: st.error("Lỗi: Không tìm thấy hồ sơ gốc (Do SĐT bị thay đổi trên Sheet?).")
                                except Exception as e: st.error(f"Lỗi: {e}")

    # 4. ADMIN
    elif st.session_state.current_page == "admin":
        st.header("⚙️ Admin"); users = sheet_users.get_all_records(); st.dataframe(users)
        with st.form("rl"):
            u = st.selectbox("User", [x['Username'] for x in users]); r = st.selectbox("Role", ["staff", "admin"])
            if st.form_submit_button("Update"): cell = sheet_users.find(u); sheet_users.update_cell(cell.row, 3, r); st.success("Done!"); st.rerun()

if st.session_state.logged_in: main_app()
else: login_screen()
